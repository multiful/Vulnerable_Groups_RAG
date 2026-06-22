# Content Hash: SHA256:TBD
# Last Updated: 2026-06-22
# Role: python-docx 기반 DOCX 파서 — RAG_PIPELINE.md §6 Parse Layer 준수
"""
born-digital DOCX (Word) 문서를 ParseIR로 변환한다.

처리 흐름:
    1. 문서 수준 순회: paragraphs + tables를 reading_order_index 순서로 순서화
    2. heading 판정: paragraph.style.name 에서 'Heading N' 추출
    3. 표: 셀 텍스트를 markdown 표로 변환 (block_type='table')
    4. 리스트: paragraph 스타일이 'List*'이면 bullet 기호 추가 후 paragraph로 처리
    5. section_path: assign_section_paths()로 사후 할당
    6. parse_hash: compute_parse_hash()로 계산
"""
from __future__ import annotations

import hashlib
import re
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

import docx
from docx.oxml.ns import qn

from ..ir_schema import ParseBlock, ParseIR
from ..cleanup import assign_section_paths, compute_parse_hash, remove_boilerplate

# 최소 텍스트 길이 (공백 제거 후)
_MIN_TEXT_LEN = 3


def slugify_docx(filename: str) -> str:
    """DOCX 파일명 → 결정론적 doc_id 슬러그."""
    stem = Path(filename).stem
    # 유니코드 정규화 후 ASCII 이외 문자는 '-'으로 치환
    normalized = unicodedata.normalize("NFKC", stem)
    slug = re.sub(r"[^\w가-힣]", "-", normalized)
    slug = re.sub(r"-{2,}", "-", slug).strip("-")
    return slug[:80]


def compute_docx_file_hash(path: Path) -> str:
    """DOCX 파일 전체 sha256. 형식: 'sha256:hex'"""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return f"sha256:{h.hexdigest()}"


def parse_with_docx(path: Path, doc_id: str | None = None, file_hash: str | None = None) -> ParseIR:
    """DOCX 파일을 파싱하여 ParseIR를 반환한다."""
    path = Path(path)
    _doc_id = doc_id or slugify_docx(path.name)
    _file_hash = file_hash or compute_docx_file_hash(path)

    document = docx.Document(str(path))
    blocks: list[ParseBlock] = []
    idx = 0

    # paragraphs와 tables를 XML 순서 기준으로 함께 순회
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            # paragraph
            para = _find_paragraph(document, child)
            if para is None:
                continue
            block = _process_paragraph(para, idx)
            if block:
                blocks.append(block)
                idx += 1

        elif tag == "tbl":
            # table
            tbl = _find_table(document, child)
            if tbl is None:
                continue
            block = _process_table(tbl, idx)
            if block:
                blocks.append(block)
                idx += 1

    blocks = remove_boilerplate(blocks)
    blocks = assign_section_paths(blocks)
    parse_hash = compute_parse_hash(blocks)

    # 페이지 수는 DOCX에서 직접 얻기 어렵다 — 단락 수 기반 추정
    estimated_pages = max(1, len(blocks) // 20)

    return ParseIR(
        doc_id=_doc_id,
        source_path=str(path),
        file_hash=_file_hash,
        parse_hash=parse_hash,
        parser_used="docx",
        page_count=estimated_pages,
        blocks=blocks,
        parse_quality_flags=[],
        parsed_at=datetime.now(timezone.utc).isoformat(),
    )


# ---------------------------------------------------------------------------
# 내부 유틸
# ---------------------------------------------------------------------------

def _find_paragraph(document, elem):
    """XML 요소에 해당하는 docx.text.paragraph.Paragraph 객체 반환."""
    for para in document.paragraphs:
        if para._element is elem:
            return para
    # 테이블 외부 단락이 document.paragraphs에 없는 경우 직접 래핑
    try:
        from docx.text.paragraph import Paragraph
        return Paragraph(elem, document)
    except Exception:
        return None


def _find_table(document, elem):
    """XML 요소에 해당하는 docx.table.Table 객체 반환."""
    for tbl in document.tables:
        if tbl._element is elem:
            return tbl
    try:
        from docx.table import Table
        return Table(elem, document)
    except Exception:
        return None


def _process_paragraph(para, idx: int) -> ParseBlock | None:
    """단락 하나를 ParseBlock으로 변환. 빈 단락이면 None 반환."""
    text = para.text.strip()
    if not text or len(text) < _MIN_TEXT_LEN:
        return None

    style_name = para.style.name if para.style else ""

    # heading 판정
    heading_match = re.match(r"[Hh]eading\s*(\d+)", style_name)
    if heading_match:
        level = int(heading_match.group(1))
        return ParseBlock(
            block_id=f"b{idx:05d}",
            block_type="heading",
            text=text,
            reading_order_index=idx,
            source_loc="docx",
            heading_level=level,
        )

    # 리스트 항목 판정 (bullet 기호 추가)
    if "List" in style_name or _has_list_format(para):
        text = f"• {text}"

    return ParseBlock(
        block_id=f"b{idx:05d}",
        block_type="paragraph",
        text=text,
        reading_order_index=idx,
        source_loc="docx",
    )


def _has_list_format(para) -> bool:
    """단락에 numPr(번호매기기) XML 요소가 있으면 리스트 항목으로 판정."""
    return para._element.find(qn("w:pPr") + "/" + qn("w:numPr")) is not None or \
           para._element.xpath("./w:pPr/w:numPr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}) != []


def _process_table(tbl, idx: int) -> ParseBlock | None:
    """표 하나를 markdown 표 텍스트로 변환한 ParseBlock 반환."""
    try:
        rows = tbl.rows
    except Exception:
        return None

    if not rows:
        return None

    md_rows: list[str] = []
    for i, row in enumerate(rows):
        cells = [c.text.strip().replace("|", "\\|").replace("\n", " ") for c in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # 헤더 구분선
            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    text = "\n".join(md_rows)
    if len(text.strip()) < _MIN_TEXT_LEN:
        return None

    return ParseBlock(
        block_id=f"b{idx:05d}",
        block_type="table",
        text=text,
        reading_order_index=idx,
        source_loc="docx:table",
    )
